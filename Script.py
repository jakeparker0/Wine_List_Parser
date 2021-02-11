import os
import operator
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK


class Wine:
    """
    Wine Object:
    The object which is created when reading the wine list
    Outputs a string in the same format for printing to the document
    """

    def __init__(self, name, year, btl, notes, region="", gls=None):
        self.name = name
        self.year = year
        self.btl = btl
        self.gls = gls
        self.region = region
        self.notes = notes

    def __str__(self):
        out = self.year + " " + self.name + " | "
        if self.region:
            out += self.region + " | "
        out += self.btl

        if self.gls:
            out += " | " + self.gls

        return out


# Docx to txt conversion


def doc_to_txt(input_dir, output_dir):
    """
    :param input_dir: Directory of docx files to be converted to txt files
    :param output_dir: Directory that the txt files will be saved into
    :return: None
    """
    for file_name in os.listdir(input_dir):
        if file_name[-5:] == ".docx" and "Wine List" in file_name and "TAKEAWAY" not in file_name:
            doc = Document(os.path.join(input_dir, file_name))
            file, ext = os.path.splitext(file_name)
            output_file = file + '.txt'
            output_path = os.path.join(output_dir, output_file)
            txt = open(output_path, "w+")
            for p in doc.paragraphs:
                txt.write(p.text + '\n')
            txt.close()


# Extracting wines from txt files
# Dictionary of all the wines
wine_dict = {}


def read_in_files(directory):
    """
    :param directory: Directory of txt files to be read into the script
    :return:
    """
    for file_name in os.listdir(directory):
        file = open(os.path.join(directory, file_name), "r")
        start(get_text(file))


def get_text(file):
    """
    :param file: txt file
    :return: List of strings of the text in the file
    """
    out = []
    for i in file.readlines(30000):
        if i != '\n':
            out.append(i)
    return out


def start(text):
    """

    :param text: Text to be read
    Will skip until it reads "Canberra Riesling, which is where the tasting notes start
    :return:
    """
    for i in range(len(text)):
        if text[i] == "Canberra Riesling \n":
            get_wines(i + 1, text)


def get_wines(pos, text):
    """
    The function looks through the text line by line and matches the pattern than the tasting notes
    are written in and creates a Wine object, which is stored in a dictionary

    :param pos: Position in text to start read at
    :param text: text to read [strings]
    :return: None
    """
    for i in range(pos, len(text)):
        line = text[i]
        if "Not wine" not in line:  # read Until this point in the document
            if "".join(list(line)[:2]) == "NV":
                year = "NV"  # no Vintage wines
            elif "".join(list(line)[:4]).isdigit():  # First 4 digits of line a number, then it is the year
                year = ("".join(list(line)[:4]))
            else:
                continue  # Otherwise skip line
            split = line.split('|')
            if len(split[0]) > 50:  # Some of the Tasting notes have a year in them and was messing it up
                continue
            # Format with region, bottle and glass
            if len(split) == 4:
                rgn, btl, gls = split[1:]

            # Other Formats
            elif len(split) == 3:
                # When no region but Glass
                if "bottle" in split[1] or "magnum" in split[1]:
                    btl, gls = split[1:]
                    rgn = ""
                else:
                    # No glass price
                    gls = None
                    rgn, btl = split[1:]

            elif len(split) == 2:
                # No region or glass
                btl = split[1]
                rgn = ""
                gls = None

            try:
                wine = Wine(split[0].split(" ", 1)[1], year, btl, text[i + 1], rgn, gls)
                # create key by removing whitespace in year and name
                key = (wine.year + wine.name.lower()).replace(" ", "")
                wine_dict.setdefault(key, wine)
                if key in wine_dict.keys():
                    if not wine_dict[key].gls and wine.gls:
                        wine_dict[key].gls = wine.gls

            except IndexError:
                print(line)
                print(split)
                continue
        else:
            break


# Output to new Document


def output_list_doc(wines):
    """
    Writes a sorted list of Wine object to a word document with formatting
    :param wines: Alphabettically sorted List of Wines
    :return: None
    """

    output_doc = Document()
    output_doc.add_heading("All Wines", 0)
    count = 0
    form = output_doc.styles['Normal'].paragraph_format
    form.keep_together = True
    form.line_spacing = 1.0

    for w in wines:
        count += 1
        p = output_doc.add_paragraph("", )
        r = p.add_run()
        f = r.font
        f.bold, f.name, f.size = True, "Galyon", Pt(11)
        r.add_text(str(w))
        r.add_break(WD_BREAK.LINE)

        n = p.add_run()
        nf = n.font
        nf.name, nf.size = "Gaylon", Pt(11)
        n.add_text(w.notes)

        # Go to next page after 11 so it looks neat
        if count == 11:
            n.add_break(WD_BREAK.PAGE)
            count = 0

    output_doc.save("output.docx")


# Run the Script

input_directory = os.path.join(os.getcwd(), "wine_lists")

txt_output_directory = os.path.join(os.getcwd(), "txt")

doc_to_txt(input_directory, txt_output_directory)

read_in_files(txt_output_directory)

output_list_doc(sorted(wine_dict.values(), key=operator.attrgetter('name')))
